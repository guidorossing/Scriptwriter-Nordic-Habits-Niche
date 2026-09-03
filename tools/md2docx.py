#!/usr/bin/env python3
"""Zet een script-.md om in een .docx voor levering.

Gebruik:  python3 tools/md2docx.py scripts/06-*.md

De 【AVATAR】-markers blijven staan én de avatar-tekst wordt vet gezet, zodat
de editor in één oogopslag ziet welke passages on-screen zijn.
"""
import os
import re
import sys

from docx import Document
from docx.shared import Pt

MARK = re.compile(r"(【AVATAR】.*?【/AVATAR】)", flags=re.S)


def convert(path):
    text = open(path, encoding="utf-8").read()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(10)

    paras = [p.strip("\n") for p in text.split("\n\n")]
    for i, para in enumerate(paras):
        if not para.strip():
            continue
        if i == 0 and para.startswith("TITLE:"):
            doc.add_heading(para[6:].strip(), level=1)
            continue
        p = doc.add_paragraph()
        for chunk in MARK.split(para):
            if not chunk:
                continue
            run = p.add_run(" ".join(chunk.split()))
            run.bold = chunk.startswith("【AVATAR】")

    out = os.path.splitext(path)[0] + ".docx"
    doc.save(out)
    print("→", out)


if __name__ == "__main__":
    for a in sys.argv[1:]:
        convert(a)
